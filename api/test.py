# from fastapi import FastAPI, HTTPException, Request, Response
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# import os
# import re
# import httpx
# import uuid
# from typing import List, Dict, Any, Optional
# import uvicorn
# from pathlib import Path
# import fitz  
# import docx
# import pandas as pd
# from langchain_openai import OpenAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_openai import ChatOpenAI
# import base64
# import traceback
# import json
# from dotenv import load_dotenv

# load_dotenv()

# app = FastAPI(title="Support Chatbot API", description="API for text and voice-based support chatbot with document retrieval")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# if not OPENAI_API_KEY:
#     print("OPENAI_API_KEY not found in environment variables")

# MODEL_ID = "gpt-4o-realtime-preview-2024-12-17"
# VOICE = "sage"
# OPENAI_SESSION_URL = "https://api.openai.com/v1/realtime/sessions"
# OPENAI_API_URL = "https://api.openai.com/v1/realtime"

# EMBEDDING_MODEL = "text-embedding-3-large"
# CHAT_MODEL = "gpt-4o"

# DOCUMENTS_DIR = Path("documents")
# VECTOR_STORE_DIR = Path("vector_store")
# CHUNK_SIZE = 1000
# CHUNK_OVERLAP = 200
# CHARS_PER_PAGE = 3000

# DOCUMENTS_DIR.mkdir(exist_ok=True)
# VECTOR_STORE_DIR.mkdir(exist_ok=True)

# vectorstore = None
# document_metadata = {}

# class TextQuery(BaseModel):
#     query: str
    
# class ChatTranscript(BaseModel):
#     question: str
#     answer: str
    
# class ChatResponse(BaseModel):
#     answer: str
#     references: List[Dict[str, Any]]
    
# class DocumentMetadata(BaseModel):
#     document_id: str
#     filename: str
#     file_type: str
#     total_pages: int
#     tables: Optional[Dict[int, int]] = None

# class DocumentsResponse(BaseModel):
#     documents: List[DocumentMetadata]

# def clean_text(text: str) -> str:
#     """Clean extracted text to improve quality and reduce noise."""
#     # Remove excessive whitespace
#     text = re.sub(r'\s+', ' ', text)
#     # Remove page numbers and footer/header artifacts
#     text = re.sub(r'\b\d+\s*of\s*\d+\b', '', text)
#     # Remove URLs that might appear in footers
#     text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
#     return text.strip()

# def extract_text_from_pdf(file_path: str) -> Dict[int, Dict]:
#     """Extract text from a PDF file maintaining original page numbers."""
#     print(f"Extracting text from PDF: {file_path}")
#     result = {}
    
#     try:
#         doc = fitz.open(file_path)
        
#         for page_num, page in enumerate(doc):
#             page = doc[page_num]
#             page_text = page.get_text("text")
#             cleaned_text = clean_text(page_text)
            
#             result[page_num] = {
#                 "text": cleaned_text,
#                 "original_page": page_num  # Store the original page number directly
#             }
                
#         doc.close()
        
#         print(f"Completed extraction from PDF: {file_path} with {len(result)} pages")
#         return result
        
#     except Exception as e:
#         print(f"Error extracting text: {str(e)}")
#         traceback.print_exc()
#         raise
    

# def extract_text_from_docx(file_path: str) -> Dict[int, Dict]:
#     """Extract text from a DOCX file with page numbering."""
#     print(f"Extracting text from DOCX: {file_path}")
#     result = {}
#     doc = docx.Document(file_path)
    
#     all_text = ""
#     for para in doc.paragraphs:
#         all_text += clean_text(para.text) + "\n"

#     total_chars = len(all_text)
#     num_pages = max(1, total_chars // CHARS_PER_PAGE + (1 if total_chars % CHARS_PER_PAGE > 0 else 0))

#     for page_num in range(num_pages):
#         start_idx = page_num * CHARS_PER_PAGE
#         end_idx = min((page_num + 1) * CHARS_PER_PAGE, total_chars)
#         page_text = all_text[start_idx:end_idx]
        
#         result[page_num] = {
#             "text": page_text,
#             "original_pages": [page_num]
#         }
    
#     print(f"Completed extraction from DOCX: {file_path} with {len(result)} standardized pages")
#     return result

# async def process_document(file_path: str, filename: str) -> str:
#     """Process a document and add it to the vector store."""
#     print(f"Processing document: {filename}")
#     file_extension = Path(filename).suffix.lower()
#     document_id = str(uuid.uuid4())
    
#     try:
#         if file_extension == '.pdf':
#             content = extract_text_from_pdf(file_path)
#             file_type = "pdf"
#         elif file_extension in ['.docx', '.doc']:
#             content = extract_text_from_docx(file_path)
#             file_type = "docx"
#         else:
#             error_msg = f"Unsupported file type: {file_extension}"
#             print(error_msg)
#             raise ValueError(error_msg)

#         document_metadata[document_id] = {
#             "filename": filename,
#             "file_type": file_type,
#             "total_pages": len(content),
#             "content": content
#         }

#         texts = []
#         metadatas = []
        
#         for page_num, page_data in content.items():
#             page_text = page_data["text"].lower()
#             text_splitter = RecursiveCharacterTextSplitter(
#                 chunk_size=CHUNK_SIZE,
#                 chunk_overlap=CHUNK_OVERLAP
#             )
#             chunks = text_splitter.split_text(page_text)
            
#             for i, chunk in enumerate(chunks):
#                 texts.append(chunk)
#                 metadatas.append({
#                     "document_id": document_id,
#                     "filename": filename,
#                     "page": page_num,
#                     "chunk": i,
#                     "source": f"{filename}, Page {page_num + 1}"
#                 })

#         print(f"Creating embeddings for {len(texts)} text chunks")
#         embeddings = OpenAIEmbeddings(model=EMBEDDING_MODEL)
        
#         global vectorstore
#         if vectorstore is None:
#             print("Creating new vector store")
#             vectorstore = FAISS.from_texts(texts=texts, embedding=embeddings, metadatas=metadatas)
#         else:
#             print("Adding to existing vector store")
#             vectorstore.add_texts(texts=texts, metadatas=metadatas)

#         print(f"Saving vector store to {VECTOR_STORE_DIR}")
#         vectorstore.save_local(str(VECTOR_STORE_DIR))
        
#         return document_id
        
#     except Exception as e:
#         print(f"Error processing document {filename}: {str(e)}")
#         traceback.print_exc()
#         raise

# def format_reference(metadata: Dict[str, Any], content: str) -> Dict[str, Any]:
#     """Format reference information to be returned to the client."""
#     document_id = metadata.get("document_id")
#     doc_metadata = document_metadata.get(document_id, {})
#     file_type = doc_metadata.get("file_type", "unknown")
#     filename = metadata.get("filename", "unknown")
#     page_number = metadata.get("page", 0)
    
#     if file_type == "unknown":
#         file_extension = filename.split('.')[-1].lower()
#         if file_extension:
#             file_type = file_extension

#     file_path = DOCUMENTS_DIR / filename

#     document_base64 = ""
#     if file_path.exists():
#         with open(file_path, "rb") as file:
#             document_binary = file.read()
#             document_base64 = base64.b64encode(document_binary).decode('utf-8')

#     return {
#         "filename": filename,
#         "file_type": file_type,
#         "file": document_base64,
#         "page": page_number,
#     }

# async def load_existing_documents():
#     """Load and process all documents from the documents directory."""
#     print("Checking for existing documents")
    
#     document_files = list(DOCUMENTS_DIR.glob("**/*"))
#     supported_extensions = ['.pdf', '.doc', '.docx']
#     document_files = [f for f in document_files if f.is_file() and f.suffix.lower() in supported_extensions]
    
#     if not document_files:
#         print("No existing documents found")
#         return
    
#     print(f"Found {len(document_files)} existing documents to process")
    
#     for file_path in document_files:
#         try:
#             await process_document(str(file_path), file_path.name)
#         except Exception as e:
#             print(f"Error processing existing document {file_path.name}: {str(e)}")

# async def load_existing_vector_store() -> None:
#     """Load existing vector store if available."""
#     global vectorstore
    
#     if VECTOR_STORE_DIR.exists() and list(VECTOR_STORE_DIR.glob("*")):
#         try:
#             print("Loading existing vector store")
#             embeddings = OpenAIEmbeddings(model=EMBEDDING_MODEL)
#             vectorstore = FAISS.load_local(str(VECTOR_STORE_DIR), embeddings, allow_dangerous_deserialization=True)
#             print("Successfully loaded existing vector store")
#         except Exception as e:
#             print(f"Error loading vector store: {str(e)}")
#             vectorstore = None
#     else:
#         print("No existing vector store found")

# @app.post("/chat", response_model=ChatResponse)
# async def text_chat(query: TextQuery):
#     """Process a text chat query and return an answer with references."""
#     print(f"Chat query received: {query.query}")
    
#     if not vectorstore:
#         error_msg = "No documents have been uploaded"
#         print(error_msg)
#         raise HTTPException(status_code=400, detail=error_msg)
    
#     try:
#         simple_messages = ["hi", "hii", "hello", "hey", "bye", "goodbye", "thanks", "thank you","kilogram","weight","dimensions","centimeters","city","street","country","handling","Understood"]
#         if query.query.lower().strip() in simple_messages:
#             print("Simple greeting/farewell detected - processing without references")
#             llm = ChatOpenAI(model_name=CHAT_MODEL, temperature=0)
#             response = llm.invoke([
#                 {"role": "system", "content": "You are a helpful support chatbot. Respond naturally to this greeting."},
#                 {"role": "user", "content": query.query}
#             ])
#             return ChatResponse(answer=response.content, references=[])

#         docs_with_scores = vectorstore.similarity_search_with_score(query.query, k=20)
#         docs = [doc for doc, score in docs_with_scores if score > 0.7]
        
#         if not docs:
#             print("No relevant documents found")
#             return ChatResponse(
#                 answer="I couldn't find any relevant information in the documents.",
#                 references=[]
#             )
        
#         print(f"Found {len(docs)} relevant documents")

#         formatted_contexts = []
#         for i, doc in enumerate(docs):
#             source = doc.metadata.get("source", f"Document {i+1}")
#             formatted_contexts.append(f"[{source}]\n{doc.page_content}")
            
#         context = "\n\n".join(formatted_contexts)
        
#         llm = ChatOpenAI(model_name=CHAT_MODEL, temperature=0)
        
#         system_prompt = """You are a helpful support chatbot. Answer the user's question based ONLY on the following context from company documents.
#         Each context section is labeled with its source in [brackets].
        
#         When answering, prefer to cite specific documents and page numbers like: "According to [Document X, Page Y]..." 
        
#         If the information is not in the context, say you don't know the provided context does not contain information about your question.
#         DO NOT make up information that is not in the context."""
        
#         messages = [
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {query.query}"}
#         ]
        
#         response = llm.invoke(messages)

#         no_info_phrases = [
#             "Sorry",
#             "I'm sorry",
#             "don't know", 
#             "couldn't find", 
#             "no information", 
#             "not mentioned", 
#             "not available",
#             "not in the context",
#             "not found in the documents",
#             "does not contain information",
#         ]
        
#         if any(phrase in response.content.lower() for phrase in no_info_phrases):
#             print("Response indicates no relevant information - returning without references")
#             return ChatResponse(answer=response.content, references=[])

#         files_dict = {}
        
#         for doc in docs:
#             filename = doc.metadata.get("filename", "Unknown")
#             if filename not in files_dict:
#                 files_dict[filename] = []
 
#             files_dict[filename].append(doc.metadata)
        
#         unique_files = {}
        
#         for doc in docs:
#             filename = doc.metadata.get("filename", "Unknown")

#             if filename not in unique_files:
#                 formatted_ref = format_reference(doc.metadata, doc.page_content)
#                 unique_files[filename] = formatted_ref

#         unique_references = list(unique_files.values())
        
#         print(f"Returning answer with {len(unique_references)} unique file references")
#         return ChatResponse(answer=response.content, references=unique_references)
    
#     except Exception as e:
#         print(f"Error in text chat: {str(e)}")
#         traceback.print_exc()
#         raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")
    
# DEFAULT_INSTRUCTIONS = """You are an expert support assistant for the company. Follow these rules:
# 1. Use only information from the company documents
# 2. If unsure, say you don't know
# 3. Reference document names and page numbers when possible
# 4. Keep your answers concise and to the point for voice interaction
# 5. If user Ask about create booking ask below question and return it as json
#     5.1. sender, receiver addresses
#     5.2. details of the package or document
#     5.3. weight and dimensions
#     5.4. Address and Shipment Details.
#     5.5. In case of an external pickup, Return toggle must be selected.
#     5.6. send this output as json to /create_booking
# """

# @app.post("/rtc-connect")
# async def connect_rtc(request: Request):
#     """Real-time WebRTC connection endpoint for voice chat with dynamic context handling."""
#     print("RTC connection request received")
#     global vectorstore
#     global document_metadata
    
#     if not vectorstore:
#         error_msg = "Please upload documents first"
#         print(error_msg)
#         raise HTTPException(status_code=400, detail=error_msg)
    
#     try:
#         client_sdp = await request.body()
#         if not client_sdp:
#             raise HTTPException(status_code=400, detail="No SDP provided")
        
#         client_sdp = client_sdp.decode()

#         # Create a document inventory instead of sending actual content
#         document_inventory = []
#         for doc_id, doc_info in document_metadata.items():
#             filename = doc_info.get("filename", "Unknown")
#             file_type = doc_info.get("file_type", "Unknown")
#             total_pages = doc_info.get("total_pages", 0)
            
#             # Create a summary of this document
#             document_inventory.append({
#                 "id": doc_id,
#                 "filename": filename,
#                 "type": file_type.upper(),
#                 "pages": total_pages
#             })
        
#         # Create a JSON representation of the document inventory
#         inventory_json = json.dumps(document_inventory)
        
#         print(f"Providing document inventory context with {len(document_inventory)} documents")
        
#         # Instructions for dynamic context retrieval
#         instructions = f"""{DEFAULT_INSTRUCTIONS}

# Important: You'll receive voice queries from users, but instead of having all document content upfront, follow this process:

# 1. Listen carefully to the user's query
# 2. Based on the query, identify what information or documents might be relevant
# 3. Respond naturally and conversationally

# Available documents in the knowledge base:
# {inventory_json}

# When responding:
# 1. If you need specific information, you can search for it in real-time
# 2. Cite specific documents when providing information
# 3. You don't need to see all document content at once - just focus on what's relevant to the current query
# 4. When you don't know something, just say so politely

# Never make up information. If you're uncertain, say you need to check the documents further."""
        
#         async with httpx.AsyncClient() as client:
#             print("Requesting ephemeral token from OpenAI")
#             token_res = await client.post(
#                 OPENAI_SESSION_URL,
#                 headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
#                 json={
#                     "model": MODEL_ID, 
#                     "modalities": ["audio", "text"],
#                     "voice": VOICE, 
#                     "input_audio_format": "pcm16",
#                     "output_audio_format": "pcm16",
#                     "input_audio_transcription": {
#                         "model": "whisper-1",
#                         "language": "en"
#                     },
#                     "tools": [
#                         {
#                             "type": "function",
#                             "function": {
#                                 "name": "search_documents",
#                                 "description": "Search the knowledge base for information based on a query",
#                                 "parameters": {
#                                     "type": "object",
#                                     "properties": {
#                                         "query": {
#                                             "type": "string",
#                                             "description": "The search query to find relevant information"
#                                         }
#                                     },
#                                     "required": ["query"]
#                                 }
#                             }
#                         }
#                     ]
#                 }
#             )
            
#             if token_res.status_code != 200:
#                 error_msg = f"Token request failed with status code {token_res.status_code}"
#                 print(error_msg)
#                 raise HTTPException(status_code=500, detail=error_msg)
            
#             token_data = token_res.json()
#             ephemeral_token = token_data.get('client_secret', {}).get('value', '')
            
#             if not ephemeral_token:
#                 error_msg = "Invalid token response"
#                 print(error_msg)
#                 raise HTTPException(status_code=500, detail=error_msg)
            
#             # Register tool handler for dynamic document searching
#             @app.post("/rtc-tool-call")
#             async def rtc_tool_call(request: Request):
#                 tool_data = await request.json()
                
#                 if tool_data.get("name") == "search_documents":
#                     query = tool_data.get("arguments", {}).get("query", "")
                    
#                     if not query:
#                         return {"result": "No query provided"}
                    
#                     # Use the vector store to find relevant information
#                     docs_with_scores = vectorstore.similarity_search_with_score(query, k=5)
#                     docs = [doc for doc, score in docs_with_scores if score > 0.7]
                    
#                     if not docs:
#                         return {"result": "No relevant information found"}
                    
#                     # Format the results in a concise way
#                     formatted_results = []
#                     for doc in docs:
#                         source = doc.metadata.get("source", "Unknown source")
#                         content = doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content
#                         formatted_results.append(f"{source}: {content}")
                    
#                     return {"result": "\n\n".join(formatted_results)}
                
#                 return {"result": "Unknown tool call"}
            
#             sdp_res = await client.post(
#                 OPENAI_API_URL,
#                 headers={
#                     "Authorization": f"Bearer {ephemeral_token}",
#                     "Content-Type": "application/sdp"
#                 },
#                 params={
#                     "model": MODEL_ID,
#                     "instructions": instructions,
#                     "voice": VOICE,
#                     "tool_url": f"{request.base_url}rtc-tool-call"
#                 },
#                 content=client_sdp
#             )
            
#             print(f"SDP exchange completed with status code {sdp_res.status_code}")
            
#             return Response(
#                 content=sdp_res.content,
#                 media_type='application/sdp',
#                 status_code=sdp_res.status_code
#             )
            
#     except Exception as e:
#         print(f"Error in RTC connection: {str(e)}")
#         traceback.print_exc()
#         raise HTTPException(status_code=500, detail=str(e))

# @app.post("/book-rtc-connect")
# async def book_connect_rtc(request: Request):
#     """Unified WebRTC endpoint with dynamic intent routing"""
#     print("Unified RTC connection request received")
#     global vectorstore
#     global document_metadata

#     # Common setup for both flows
#     client_sdp = await request.body()
#     if not client_sdp:
#         raise HTTPException(status_code=400, detail="No SDP provided")
    
#     client_sdp = client_sdp.decode()

#     # Dynamic instructions template
#     base_instructions = """Respond based on user intent:
#     - If user mentions booking, shipping, or shipment creation: Use booking workflow
#     - Else: Use general knowledge base
#     ---"""

#     # Prepare context for both scenarios
#     context = []
#     if vectorstore and document_metadata:
#         context = [f"Document: {doc_info.get('filename')} ({doc_info.get('file_type')})" 
#                   for doc_id, doc_info in document_metadata.items()]
    
#     full_instructions = f"""{base_instructions}
    
#     {DEFAULT_INSTRUCTIONS}
    
#     {'## Available Documents ##' if context else ''}
#     {chr(10).join(context)}
    
#     ## Booking Workflow ##
#         1.Collect information in this exact order:\n
#             -Full sender address (street, city, country)\n
#             -Full receiver address (street, city, country)\n
#             -Package or document description\n
#             -Weight in kilograms\n
#             -Dimensions (length × width × height in cm)\n
#             -Special handling requirements\n
#         2.Return pickup needed (yes/no)
#         3.Ask one question at a time.
#         4.Do not confirm each detail before proceeding to the next.
#         5.After collecting all required information:
#             -Generate a single 6-digit booking number (only once, and reuse it throughout the process).
#             -Display all collected details, each on a new line. Begin with the booking number. Each new detail must start on a new line in the order above.
#             like this format:
#                 -Booking number : Booking number \n
#                 -Render address :Full sender address \n
#                 -Receiver address \n
#                 -Package or document description \n
#                 -Weight in kilograms \n
#                 -Dimensions (length × width × height in cm) \n
#                 -Special handling requirements \n
#                 -Ask the user for final confirmation.
#         Only proceed to create the booking after the user gives final confirmation.
#         If the user wants to correct a detail, handle the correction gracefully and update only the specified part.
#         **only Booking number should come after confirmation not other**"""
    
#     try:
#         async with httpx.AsyncClient() as client:
#             # Get session token
#             token_res = await client.post(
#                 OPENAI_SESSION_URL,
#                 headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
#                 json={
#                     "model": MODEL_ID, 
#                     "modalities": ["audio", "text"],
#                     "voice": VOICE, 
#                     "input_audio_format": "pcm16",
#                     "output_audio_format": "pcm16",
#                     "input_audio_transcription": {
#                         "model": "whisper-1",
#                         "language": "en"
#                     },
#                 }
#             )

#             if token_res.status_code != 200:
#                 raise HTTPException(status_code=500, detail="Token request failed")

#             token_data = token_res.json()
#             ephemeral_token = token_data.get('client_secret', {}).get('value', '')

#             # Dynamic parameter configuration
#             params = {
#                 "model": MODEL_ID,
#                 "instructions": full_instructions,
#                 "voice": VOICE,
#                 "context": {
#                     "auto_route": True,
#                     "booking_keywords": ["book", "shipment", "package", "create booking"],
#                     "general_keywords": ["help", "question", "support"]
#                 }
#             }

#             # Single SDP exchange with dynamic instructions
#             sdp_res = await client.post(
#                 OPENAI_API_URL,
#                 headers={
#                     "Authorization": f"Bearer {ephemeral_token}",
#                     "Content-Type": "application/sdp"
#                 },
#                 params=params,
#                 content=client_sdp
#             )

#             return Response(
#                 content=sdp_res.content,
#                 media_type='application/sdp',
#                 status_code=sdp_res.status_code
#             )

#     except Exception as e:
#         return Response(
#             status_code=500,
#             content={"detail": f"Unified assistant error: {str(e)}"}
#         )     
     
# @app.post("/references_for_query", response_model=ChatResponse)
# async def references_for_query(transcript: ChatTranscript):
#     """Retrieve references for a given text query."""
#     print(f"Query received: {transcript.question}")
#     print(f"Response received: {transcript.answer}")
    
#     if not vectorstore:
#         error_msg = "No documents have been uploaded"
#         print(error_msg)
#         raise HTTPException(status_code=400, detail=error_msg)
    
#     try:
#         random_phrases = [
#             "hii", "hi", "hello", "hey", "bye", "goodbye", "thanks", "thank you","book","product","yes","um","address","package","return","kilogram","weight","dimensions","centimeters","city","street","country","handling","Understood"
#         ]
        
#         if any(phrase in transcript.question.lower() for phrase in random_phrases):
#             print("Response indicates no relevant information - returning without references")
#             return ChatResponse(answer=transcript.question, references=[])
        
#         no_info_phrases = [
#             "don't know", 
#             "couldn't find", 
#             "no information", 
#             "not mentioned", 
#             "not available",
#             "not in the context",
#             "not found in the documents",
#             "does not contain information",
#             "book","product","yes","um","address","package","return","kilogram","weight","dimensions","centimeters","city","street","country","handling","Understood"
#         ]
        
#         if any(phrase in transcript.answer.lower() for phrase in no_info_phrases):
#             print("Response indicates no relevant information - returning without references")
#             return ChatResponse(answer=transcript.question, references=[])
        
#         docs_with_scores = vectorstore.similarity_search_with_score(transcript.question, k=10)
#         docs = [doc for doc, score in docs_with_scores if score > 0.7]
        
#         if not docs:
#             print("No relevant documents found")
#             return ChatResponse(answer=transcript.question, references=[])
        
#         print(f"Found {len(docs)} relevant documents")
        
#         unique_files = {}
        
#         for doc in docs:
#             filename = doc.metadata.get("filename", "Unknown")

#             if filename not in unique_files:
#                 formatted_ref = format_reference(doc.metadata, doc.page_content)
#                 unique_files[filename] = formatted_ref

#         unique_references = list(unique_files.values())
        
#         return ChatResponse(answer=transcript.question, references=unique_references)
    
#     except Exception as e:
#         print(f"Error in references_for_query: {str(e)}")
#         traceback.print_exc()
#         raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")

# @app.on_event("startup")
# async def startup_event():
#     """Load existing vector store and documents on startup."""
#     print("Starting application...")
#     await load_existing_vector_store()
#     await load_existing_documents()
#     print(f"Startup complete. Vector store initialized: {vectorstore is not None}")
#     print(f"Total documents loaded: {len(document_metadata)}")
    

# if __name__ == "__main__":
#     uvicorn.run("main:app", host="127.0.0.1", port=8001, reload=True)

# from fastapi import FastAPI, HTTPException, Request, Response, WebSocket, WebSocketDisconnect, Depends
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# import os
# import httpx
# import uuid
# from typing import List, Dict, Any, Optional, Set
# import uvicorn
# from pathlib import Path
# import fitz 
# import docx
# import pandas as pd
# from langchain_openai import OpenAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_openai import ChatOpenAI
# import base64
# import json
# from dotenv import load_dotenv

# load_dotenv()

# app = FastAPI(title="Support Chatbot API", description="API for text and voice-based support chatbot with document retrieval")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# if not OPENAI_API_KEY:
#     print("OPENAI_API_KEY not found in environment variables")

# MODEL_ID = "gpt-4o-realtime-preview-2024-12-17"
# VOICE = "sage"
# OPENAI_SESSION_URL = "https://api.openai.com/v1/realtime/sessions"
# OPENAI_API_URL = "https://api.openai.com/v1/realtime"

# EMBEDDING_MODEL = "text-embedding-3-large"
# CHAT_MODEL = "gpt-4o"

# DOCUMENTS_DIR = Path("documents")
# VECTOR_STORE_DIR = Path("vector_store")
# CHUNK_SIZE = 1000
# CHUNK_OVERLAP = 200

# DOCUMENTS_DIR.mkdir(exist_ok=True)
# VECTOR_STORE_DIR.mkdir(exist_ok=True)

# vectorstore = None
# document_metadata = {}
# session_contexts: Dict[str, Dict[str, Any]] = {}

# class TextQuery(BaseModel):
#     query: str
    
# class ChatResponse(BaseModel):
#     answer: str
#     references: List[Dict[str, Any]]
    
# class DocumentMetadata(BaseModel):
#     document_id: str
#     filename: str
#     file_type: str
#     total_pages: int
#     tables: Optional[Dict[int, int]] = None

# class DocumentsResponse(BaseModel):
#     documents: List[DocumentMetadata]
    
# class WebRTCSession(BaseModel):
#     session_id: str
#     query: Optional[str] = None

# class QueryTranscript(BaseModel):
#     session_id: str
#     transcript: str
    
# class ConnectionManager:
#     def __init__(self):
#         self.active_connections: Dict[str, WebSocket] = {}
#         self.session_queries: Dict[str, Set[str]] = {}
        
#     async def connect(self, websocket: WebSocket, session_id: str):
#         await websocket.accept()
#         self.active_connections[session_id] = websocket
#         self.session_queries[session_id] = set()
#         print(f"New WebSocket connection: {session_id}")
        
#     def disconnect(self, session_id: str):
#         if session_id in self.active_connections:
#             del self.active_connections[session_id]
#         if session_id in self.session_queries:
#             del self.session_queries[session_id]
#         print(f"WebSocket disconnected: {session_id}")
        
#     async def send_document_update(self, session_id: str, documents: List[Dict[str, Any]]):
#         if session_id in self.active_connections:
#             await self.active_connections[session_id].send_json({
#                 "type": "document_update",
#                 "documents": documents
#             })
            
#     def add_query(self, session_id: str, query: str):
#         if session_id in self.session_queries:
#             self.session_queries[session_id].add(query)
            
#     def get_queries(self, session_id: str) -> List[str]:
#         return list(self.session_queries.get(session_id, set()))
    
# manager = ConnectionManager()

# def extract_text_and_tables_from_pdf(file_path: str) -> Dict[int, Dict]:
#     """Extract text, tables, and images from a PDF file."""
#     print(f"Extracting content from PDF: {file_path}")
#     result = {}
#     doc = fitz.open(file_path)
    
#     for page_num, page in enumerate(doc):
#         page_text = page.get_text()
#         page_tables = []
#         tables = page.find_tables()
#         if tables and tables.tables:
#             for table_idx, table in enumerate(tables.tables):
#                 df = pd.DataFrame([[str(cell) if hasattr(cell, 'text') else str(cell) for cell in row] for row in table.cells])
#                 page_tables.append(df.to_dict())

#         page_images = []
#         images = page.get_images(full=True)
#         for img_index, img in enumerate(images):
#             xref = img[0]
#             base_image = doc.extract_image(xref)
#             image_data = base_image["image"]
#             image_b64 = base64.b64encode(image_data).decode('utf-8')
#             page_images.append({
#                 "image_id": f"img_{page_num}_{img_index}",
#                 "image_data": image_b64,
#                 "width": base_image["width"],
#                 "height": base_image["height"]
#             })
        
#         result[page_num] = {
#             "text": page_text,
#             "tables": page_tables,
#             "images": page_images
#         }
    
#     doc.close()
#     print(f"Completed extraction from PDF: {file_path}")
#     return result

# def extract_text_and_tables_from_docx(file_path: str) -> Dict[int, Dict]:
#     """Extract text, tables, and images from a DOCX file with improved page detection."""
#     print(f"Extracting content from DOCX: {file_path}")
#     result = {}
#     doc = docx.Document(file_path)

#     CHARS_PER_PAGE = 3000
    
#     all_text = ""
#     for para in doc.paragraphs:
#         all_text += para.text + "\n"

#     total_chars = len(all_text)
#     num_pages = max(1, total_chars // CHARS_PER_PAGE + (1 if total_chars % CHARS_PER_PAGE > 0 else 0))

#     for page_num in range(num_pages):
#         start_idx = page_num * CHARS_PER_PAGE
#         end_idx = min((page_num + 1) * CHARS_PER_PAGE, total_chars)
#         page_text = all_text[start_idx:end_idx]

#         tables_for_page = []
#         total_tables = len(doc.tables)
#         tables_start_idx = (page_num * total_tables) // num_pages
#         tables_end_idx = ((page_num + 1) * total_tables) // num_pages
        
#         for table_idx in range(tables_start_idx, tables_end_idx):
#             if table_idx < total_tables:
#                 table = doc.tables[table_idx]
#                 table_data = []
#                 for row in table.rows:
#                     row_data = [cell.text for cell in row.cells]
#                     table_data.append(row_data)
#                 tables_for_page.append(pd.DataFrame(table_data).to_dict())
        
#         result[page_num] = {
#             "text": page_text,
#             "tables": tables_for_page,
#             "images": []
#         }
    
#     print(f"Completed extraction from DOCX: {file_path}")
#     return result

# async def process_document(file_path: str, filename: str) -> str:
#     """Process a document and add it to the vector store."""
#     print(f"Processing document: {filename}")
#     file_extension = Path(filename).suffix.lower()
#     document_id = str(uuid.uuid4())
    
#     try:
#         if file_extension == '.pdf':
#             content = extract_text_and_tables_from_pdf(file_path)
#             file_type = "pdf"
#         elif file_extension in ['.docx', '.doc']:
#             content = extract_text_and_tables_from_docx(file_path)
#             file_type = "docx"
#         else:
#             error_msg = f"Unsupported file type: {file_extension}"
#             print(error_msg)
#             raise ValueError(error_msg)

#         document_metadata[document_id] = {
#             "filename": filename,
#             "file_type": file_type,
#             "total_pages": len(content),
#             "content": content,
#             "tables": {page_num: len(page_data.get("tables", [])) for page_num, page_data in content.items()}
#         }

#         texts = []
#         metadatas = []
        
#         for page_num, page_data in content.items():
#             page_text = page_data["text"]
#             text_splitter = RecursiveCharacterTextSplitter(
#                 chunk_size=CHUNK_SIZE,
#                 chunk_overlap=CHUNK_OVERLAP
#             )
#             chunks = text_splitter.split_text(page_text)
            
#             for i, chunk in enumerate(chunks):
#                 texts.append(chunk)
#                 metadatas.append({
#                     "document_id": document_id,
#                     "filename": filename,
#                     "page": page_num,
#                     "chunk": i,
#                     "source": f"{filename}, Page {page_num + 1}"
#                 })

#         print(f"Creating embeddings for {len(texts)} text chunks")
#         embeddings = OpenAIEmbeddings(model=EMBEDDING_MODEL)
        
#         global vectorstore
#         if vectorstore is None:
#             print("Creating new vector store")
#             vectorstore = FAISS.from_texts(texts=texts, embedding=embeddings, metadatas=metadatas)
#         else:
#             print("Adding to existing vector store")
#             vectorstore.add_texts(texts=texts, metadatas=metadatas)

#         print(f"Saving vector store to {VECTOR_STORE_DIR}")
#         vectorstore.save_local(str(VECTOR_STORE_DIR))
        
#         return document_id
        
#     except Exception as e:
#         print(f"Error processing document {filename}: {str(e)}")
#         raise

# def format_reference(metadata: Dict[str, Any], content: str) -> Dict[str, Any]:
#     """Format reference information to be returned to the client."""
#     document_id = metadata.get("document_id")
#     doc_metadata = document_metadata.get(document_id, {})
#     file_type = doc_metadata.get("file_type", "unknown")
#     filename = metadata.get("filename", "unknown")
    
#     if file_type == "unknown":
#         file_extension = filename.split('.')[-1].lower()
#         if file_extension:
#             file_type = file_extension

#     file_path = DOCUMENTS_DIR / filename

#     document_base64 = ""
#     if file_path.exists():
#         with open(file_path, "rb") as file:
#             document_binary = file.read()
#             document_base64 = base64.b64encode(document_binary).decode('utf-8')

#     return {
#         "filename": filename,
#         "file_type": file_type,
#         "file": document_base64,
#     }

# async def load_existing_documents():
#     """Load and process all documents from the documents directory."""
#     print("Checking for existing documents")
    
#     document_files = list(DOCUMENTS_DIR.glob("**/*"))
#     supported_extensions = ['.pdf', '.doc', '.docx']
#     document_files = [f for f in document_files if f.is_file() and f.suffix.lower() in supported_extensions]
    
#     if not document_files:
#         print("No existing documents found")
#         return
    
#     print(f"Found {len(document_files)} existing documents to process")
    
#     for file_path in document_files:
#         try:
#             await process_document(str(file_path), file_path.name)
#         except Exception as e:
#             print(f"Error processing existing document {file_path.name}: {str(e)}")

# async def load_existing_vector_store() -> None:
#     """Load existing vector store if available."""
#     global vectorstore
    
#     if VECTOR_STORE_DIR.exists() and list(VECTOR_STORE_DIR.glob("*")):
#         try:
#             print("Loading existing vector store")
#             embeddings = OpenAIEmbeddings(model=EMBEDDING_MODEL)
#             vectorstore = FAISS.load_local(str(VECTOR_STORE_DIR), embeddings, allow_dangerous_deserialization=True)
#             print("Successfully loaded existing vector store")
#         except Exception as e:
#             print(f"Error loading vector store: {str(e)}")
#             vectorstore = None
#     else:
#         print("No existing vector store found")
        
# def update_session_context(session_id: str, query: str, relevant_docs: List[Dict[str, Any]]):
#     if session_id not in session_contexts:
#         session_contexts[session_id] = {
#             "context": "",
#             "instructions": DEFAULT_INSTRUCTIONS
#         }

#     context = "\n".join([f"Document: {doc['filename']}, Page: {doc['page']}\n{doc['content']}" for doc in relevant_docs])
#     session_contexts[session_id]["context"] = context

# def get_session_context(session_id: str) -> Optional[Dict[str, Any]]:
#     return session_contexts.get(session_id)


# @app.post("/chat", response_model=ChatResponse)
# async def text_chat(query: TextQuery):
#     """Process a text chat query and return an answer with references."""
#     print(f"Chat query received: {query.query}")
    
#     if not vectorstore:
#         error_msg = "No documents have been uploaded"
#         print(error_msg)
#         raise HTTPException(status_code=400, detail=error_msg)
    
#     try:
#         docs = vectorstore.similarity_search(query.query, k=4)
        
#         if not docs:
#             print("No relevant documents found")
#             return ChatResponse(
#                 answer="I couldn't find any relevant information in the documents.",
#                 references=[]
#             )
        
#         print(f"Found {len(docs)} relevant documents")
        
#         context = "\n\n".join([doc.page_content for doc in docs])
        
#         llm = ChatOpenAI(model_name=CHAT_MODEL, temperature=0)
        
#         system_prompt = """You are a helpful support chatbot. Answer the user's question based ONLY on the following context from company documents.
#         If the information is not in the context, say you don't know. Reference document names and page numbers when possible.

#         DO NOT make up information that is not in the context and don't use bold and ** characters for reply."""
        
#         messages = [
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {query.query}"}
#         ]
        
#         response = llm.invoke(messages)

#         unique_files = {}
        
#         for doc in docs:
#             filename = doc.metadata.get("filename", "Unknown")

#             if filename not in unique_files:
#                 formatted_ref = format_reference(doc.metadata, doc.page_content)
#                 unique_files[filename] = formatted_ref

#         unique_references = list(unique_files.values())
        
#         print(f"Returning answer with {len(unique_references)} unique file references")
#         return ChatResponse(answer=response.content, references=unique_references)
    
#     except Exception as e:
#         print(f"Error processing query: {str(e)}")
#         raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")
    
# DEFAULT_INSTRUCTIONS = """You are an expert support assistant for the company. Follow these rules:
# 1. Use only information from the company documents
# 2. If unsure, say you don't know
# 3. Reference document names and page numbers when possible
# 4. Keep your answers concise and to the point for voice interaction"""

# @app.websocket("/ws/{session_id}")
# async def websocket_endpoint(websocket: WebSocket, session_id: str):
#     """WebSocket endpoint for real-time document updates during voice chat."""
#     try:
#         await manager.connect(websocket, session_id)
#         while True:
#             data = await websocket.receive_text()
#             try:
#                 message = json.loads(data)
#                 if message.get("type") == "query":
#                     query = message.get("query", "")
#                     if query and len(query) > 3:
#                         if vectorstore:
#                             docs = vectorstore.similarity_search(query, k=3)
#                             documents = []
#                             for doc in docs:
#                                 documents.append({
#                                     "filename": doc.metadata.get("filename", "Unknown"),
#                                     "page": doc.metadata.get("page", 0) + 1,
#                                     "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
#                                     "source": f"{doc.metadata.get('filename', 'Unknown')}, Page {doc.metadata.get('page', 0) + 1}"
#                                 })
              
#                             await manager.send_document_update(session_id, documents)
#                             manager.add_query(session_id, query)
#             except json.JSONDecodeError:
#                 pass
#     except WebSocketDisconnect:
#         manager.disconnect(session_id)

# @app.post("/rtc/start-session")
# async def start_rtc_session(session: WebRTCSession):
#     """Start a new WebRTC session with optional initial query."""
#     if not session.session_id:
#         session.session_id = str(uuid.uuid4())
        
#     initial_context = ""
#     if session.query and vectorstore:
#         docs = vectorstore.similarity_search(session.query, k=3)
#         initial_context = "\n".join([
#             f"Document: {doc.metadata.get('filename', 'Unknown')}, "
#             f"Page: {doc.metadata.get('page', 0) + 1}\n"
#             f"{doc.page_content}\n" 
#             for doc in docs
#         ])
        
#     return {
#         "session_id": session.session_id,
#         "initial_context": initial_context
#     }

# @app.post("/rtc/process-transcript")
# async def process_transcript(query_data: QueryTranscript):
#     """Process a transcript from the voice chat to extract queries and get relevant documents."""
#     if not query_data.session_id or not query_data.transcript:
#         return {"error": "Missing session_id or transcript"}
    
#     try:
#         llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0)
        
#         extract_query_prompt = """
#         Extract the main question or query from this transcript of a voice conversation.
#         Return only the most important question or information request.
#         If there are multiple questions, focus on the most recent or most significant one.
        
#         Transcript: {transcript}
#         """
        
#         messages = [
#             {"role": "system", "content": extract_query_prompt.format(transcript=query_data.transcript)}
#         ]
        
#         response = llm.invoke(messages)
#         extracted_query = response.content.strip()
        
#         if vectorstore and extracted_query:
#             docs = vectorstore.similarity_search(extracted_query, k=3)
            
#             documents = []
#             for doc in docs:
#                 documents.append({
#                     "filename": doc.metadata.get("filename", "Unknown"),
#                     "page": doc.metadata.get("page", 0) + 1,
#                     "content": doc.page_content,
#                     "source": f"{doc.metadata.get('filename', 'Unknown')}, Page {doc.metadata.get('page', 0) + 1}"
#                 })

#             # Update session context with relevant documents
#             update_session_context(query_data.session_id, extracted_query, documents)

#             if query_data.session_id in manager.active_connections:
#                 await manager.send_document_update(query_data.session_id, documents)
 
#             manager.add_query(query_data.session_id, extracted_query)
            
#             return {
#                 "query": extracted_query,
#                 "documents": documents
#             }
        
#         return {"query": extracted_query, "documents": []}
        
#     except Exception as e:
#         print(f"Error processing transcript: {str(e)}")
#         return {"error": str(e)}
    
# @app.post("/rtc/update-instructions/{session_id}")
# async def update_instructions(session_id: str, new_instructions: str):
#     """Update the instructions for a specific session."""
#     if not session_id or not new_instructions:
#         return {"error": "Missing session_id or new_instructions"}
    
#     if session_id in session_contexts:
#         session_contexts[session_id]["instructions"] = new_instructions
#         return {"status": "Instructions updated successfully"}
#     else:
#         return {"error": "Session not found"}

# @app.get("/rtc/session-history/{session_id}")
# async def get_session_history(session_id: str):
#     """Get the history of queries for a specific session."""
#     if not session_id:
#         return {"error": "Missing session_id"}
    
#     queries = manager.get_queries(session_id)
    
#     return {
#         "session_id": session_id,
#         "queries": queries
#     }

# @app.post("/book-rtc-connect")
# async def connect_rtc(request: Request):
#     """Enhanced real-time WebRTC connection endpoint for voice chat."""
#     print("RTC connection request received")
#     global vectorstore
    
#     if not vectorstore:
#         error_msg = "Please upload documents first"
#         print(error_msg)
#         raise HTTPException(status_code=400, detail=error_msg)
    
#     try:
#         body = await request.body()
#         body_dict = {}
        
#         try:
#             body_str = body.decode()
#             if body_str.startswith("{"):
#                 body_dict = json.loads(body_str)
#                 client_sdp = body_dict.get("sdp", "")
#                 session_id = body_dict.get("session_id", str(uuid.uuid4()))
#                 initial_query = body_dict.get("initial_query", "")
#             else:
#                 client_sdp = body_str
#                 session_id = str(uuid.uuid4())
#                 initial_query = ""
#         except:
#             client_sdp = body.decode()
#             session_id = str(uuid.uuid4())
#             initial_query = ""
        
#         if not client_sdp:
#             raise HTTPException(status_code=400, detail="No SDP provided")

#         # Initialize session context
#         session_contexts[session_id] = {
#             "context": "",
#             "instructions": DEFAULT_INSTRUCTIONS
#         }

#         instructions = DEFAULT_INSTRUCTIONS
        
#         async with httpx.AsyncClient() as client:
#             print("Requesting ephemeral token from OpenAI")
#             token_res = await client.post(
#                 OPENAI_SESSION_URL,
#                 headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
#                 json={
#                     "model": MODEL_ID, 
#                     "modalities": ["audio", "text"],
#                     "voice": VOICE, 
#                     "input_audio_format": "pcm16",
#                     "output_audio_format": "pcm16",
#                     "input_audio_transcription": {
#                         "model": "whisper-1",
#                         "language": "en"
#                     },
#                 }
#             )
            
#             if token_res.status_code != 200:
#                 error_msg = f"Token request failed with status code {token_res.status_code}"
#                 print(error_msg)
#                 raise HTTPException(status_code=500, detail=error_msg)
            
#             token_data = token_res.json()
#             ephemeral_token = token_data.get('client_secret', {}).get('value', '')
            
#             if not ephemeral_token:
#                 error_msg = "Invalid token response"
#                 print(error_msg)
#                 raise HTTPException(status_code=500, detail=error_msg)
            
#             sdp_res = await client.post(
#                 OPENAI_API_URL,
#                 headers={
#                     "Authorization": f"Bearer {ephemeral_token}",
#                     "Content-Type": "application/sdp"
#                 },
#                 params={
#                     "model": MODEL_ID,
#                     "instructions": instructions,
#                     "voice": VOICE,
#                     "session_id": session_id,
#                 },
#                 content=client_sdp
#             )
            
#             print(f"SDP exchange completed with status code {sdp_res.status_code}")
 
#             response = Response(
#                 content=sdp_res.content,
#                 media_type='application/sdp',
#                 status_code=sdp_res.status_code,
#                 headers={"X-Session-ID": session_id}
#             )
            
#             return response
            
#     except Exception as e:
#         print(f"Error in RTC connection: {str(e)}", exc_info=True)
#         raise HTTPException(status_code=500, detail=str(e))

# @app.on_event("startup")
# async def startup_event():
#     """Load existing vector store and documents on startup."""
#     print("Starting application...")
#     await load_existing_vector_store()
#     await load_existing_documents()
#     print(f"Startup complete. Vector store initialized: {vectorstore is not None}")
#     print(f"Total documents loaded: {len(document_metadata)}")
    

# if __name__ == "__main__":
#     uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)


from fastapi import FastAPI, HTTPException, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os
import re
import httpx
import uuid
from typing import List, Dict, Any, Optional
import uvicorn
from pathlib import Path
import fitz  
import docx
import pandas as pd
import base64
import traceback
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

app = FastAPI(title="Support Chatbot API", description="API for text and voice-based support chatbot with document retrieval")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    print("OPENAI_API_KEY not found in environment variables")

# OpenAI client initialization
client = OpenAI(api_key=OPENAI_API_KEY)

MODEL_ID = "gpt-4o-realtime-preview-2024-12-17"
VOICE = "sage"
OPENAI_SESSION_URL = "https://api.openai.com/v1/realtime/sessions"
OPENAI_API_URL = "https://api.openai.com/v1/realtime"

CHAT_MODEL = "gpt-4o"

DOCUMENTS_DIR = Path("documents")
VECTOR_STORE_ID = "vs_682d690a20d48191aef3b1fd3cc5cbf3"  # Your vector store ID
CHARS_PER_PAGE = 3000

DOCUMENTS_DIR.mkdir(exist_ok=True)

document_metadata = {}

class TextQuery(BaseModel):
    query: str
    
class ChatTranscript(BaseModel):
    question: str
    answer: str
    
class ChatResponse(BaseModel):
    answer: str
    references: List[Dict[str, Any]]
    
class DocumentMetadata(BaseModel):
    document_id: str
    filename: str
    file_type: str
    total_pages: int
    tables: Optional[Dict[int, int]] = None

class DocumentsResponse(BaseModel):
    documents: List[DocumentMetadata]

def clean_text(text: str) -> str:
    """Clean extracted text to improve quality and reduce noise."""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove common header/footer patterns like "King & Wood Mallesons"
    text = re.sub(r'King\s*&\s*Wood\s*Mallesons.*?(?:Version|Date):\s*[\d\.]+.*?(?:\d{2}/\d{2}/\d{4})', '', text, flags=re.IGNORECASE)
    
    # Remove page numbers
    text = re.sub(r'\b\d+\s*of\s*\d+\b', '', text)
    
    # Remove "User Guide" and similar document identifiers
    text = re.sub(r'User\s+Guide', '', text, flags=re.IGNORECASE)
    
    # Remove Digital Hub Mobile App and related text
    text = re.sub(r'Digital\s+Hub\s+Mobile\s+App.*?Setting\s+Up', '', text, flags=re.IGNORECASE)
    
    # Remove image descriptions (often in brackets or with "Figure" prefix)
    text = re.sub(r'\[.*?\]', '', text)  # Text in square brackets
    text = re.sub(r'Figure\s+\d+:.*?\.', '', text)  # Figure captions
    
    # Remove URLs that might appear in footers
    text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
    
    return text.strip()

def extract_text_from_pdf(file_path: str) -> Dict[int, Dict]:
    """Extract text from a PDF file maintaining original page numbers while removing headers/footers."""
    print(f"Extracting text from PDF: {file_path}")
    result = {}
    
    try:
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            page = doc[page_num]
            
            # Get text blocks to analyze layout
            blocks = page.get_text("blocks")
            
            # Skip potential headers (top 10% of page) and footers (bottom 10% of page)
            filtered_blocks = []
            page_height = page.rect.height
            
            for block in blocks:
                y_pos = block[1]  # Top y-coordinate of the block
                
                # Skip blocks that are likely headers or footers
                if y_pos < page_height * 0.1 or y_pos > page_height * 0.9:
                    # Check if the block contains header/footer patterns
                    if re.search(r'King\s*&\s*Wood|Version|User Guide|Digital\s+Hub|Page \d+|\d+ of \d+', block[4]):
                        continue
                
                filtered_blocks.append(block[4])  # Add block text
            
            # Join the filtered content
            page_text = " ".join(filtered_blocks)
            cleaned_text = clean_text(page_text)
            
            # Only store pages with meaningful content
            if len(cleaned_text) > 5:  # Minimum content threshold
                result[page_num] = {
                    "text": cleaned_text,
                    "original_page": page_num
                }
                
        doc.close()
        
        print(f"Completed extraction from PDF: {file_path} with {len(result)} pages")
        return result
        
    except Exception as e:
        print(f"Error extracting text: {str(e)}")
        traceback.print_exc()
        raise

def extract_text_from_docx(file_path: str) -> Dict[int, Dict]:
    """Extract text from a DOCX file with page numbering."""
    print(f"Extracting text from DOCX: {file_path}")
    result = {}
    doc = docx.Document(file_path)
    
    all_text = ""
    for para in doc.paragraphs:
        all_text += clean_text(para.text) + "\n"

    total_chars = len(all_text)
    num_pages = max(1, total_chars // CHARS_PER_PAGE + (1 if total_chars % CHARS_PER_PAGE > 0 else 0))

    for page_num in range(num_pages):
        start_idx = page_num * CHARS_PER_PAGE
        end_idx = min((page_num + 1) * CHARS_PER_PAGE, total_chars)
        page_text = all_text[start_idx:end_idx]
        
        result[page_num] = {
            "text": page_text,
            "original_pages": [page_num]
        }
    
    print(f"Completed extraction from DOCX: {file_path} with {len(result)} standardized pages")
    return result

async def process_document(file_path: str, filename: str) -> str:
    """Process a document and save metadata."""
    print(f"Processing document: {filename}")
    file_extension = Path(filename).suffix.lower()
    document_id = str(uuid.uuid4())
    
    try:
        if file_extension == '.pdf':
            content = extract_text_from_pdf(file_path)
            file_type = "pdf"
        elif file_extension in ['.docx', '.doc']:
            content = extract_text_from_docx(file_path)
            file_type = "docx"
        else:
            error_msg = f"Unsupported file type: {file_extension}"
            print(error_msg)
            raise ValueError(error_msg)

        document_metadata[document_id] = {
            "filename": filename,
            "file_type": file_type,
            "total_pages": len(content),
            "content": content
        }
        
        print(f"Document processed: {filename}")
        
        # Note: Vector store handling now managed by OpenAI through file uploads
        
        return document_id
        
    except Exception as e:
        print(f"Error processing document {filename}: {str(e)}")
        traceback.print_exc()
        raise

def format_reference(file_data: Dict[str, Any]) -> Dict[str, Any]:
    """Format reference information from file citation data."""
    filename = file_data.get("filename", "unknown")
    file_id = file_data.get("file_id", "")
    
    # Determine file type from filename
    file_extension = filename.split('.')[-1].lower() if '.' in filename else "unknown"
    
    file_path = DOCUMENTS_DIR / filename

    document_base64 = ""
    if file_path.exists():
        with open(file_path, "rb") as file:
            document_binary = file.read()
            document_base64 = base64.b64encode(document_binary).decode('utf-8')

    return {
        "filename": filename,
        "file_type": file_extension,
        "file": document_base64,
        "file_id": file_id,
        "page": 0,  # Default page if not available
    }

async def load_existing_documents():
    """Load document metadata from the documents directory."""
    print("Checking for existing documents")
    
    document_files = list(DOCUMENTS_DIR.glob("**/*"))
    supported_extensions = ['.pdf', '.doc', '.docx']
    document_files = [f for f in document_files if f.is_file() and f.suffix.lower() in supported_extensions]
    
    if not document_files:
        print("No existing documents found")
        return
    
    print(f"Found {len(document_files)} existing documents to process")
    
    for file_path in document_files:
        try:
            await process_document(str(file_path), file_path.name)
        except Exception as e:
            print(f"Error processing existing document {file_path.name}: {str(e)}")

@app.post("/chat", response_model=ChatResponse)
async def text_chat(query: TextQuery):
    """Process a text chat query and return an answer with references."""
    print(f"Chat query received: {query.query}")
    
    try:
        simple_messages = ["hi", "hii", "hello", "hey", "bye", "goodbye", "thanks", "thank you", "kilogram", "weight", "dimensions", "centimeters", "city", "street", "country", "handling", "Understood"]
        if query.query.lower().strip() in simple_messages:
            print("Simple greeting/farewell detected - processing without references")
            response = client.chat.completions.create(
                model=CHAT_MODEL,
                messages=[
                    {"role": "system", "content": "You are a helpful support chatbot. Respond naturally to this greeting."},
                    {"role": "user", "content": query.query}
                ],
                temperature=0
            )
            return ChatResponse(answer=response.choices[0].message.content, references=[])

        # Use OpenAI's file search tool
        response = client.responses.create(
            model=CHAT_MODEL,
            input=f"Answer this question based only on the information in the uploaded documents: {query.query}",
            tools=[{
                "type": "file_search",
                "vector_store_ids": [VECTOR_STORE_ID]
            }]
        )
        
        # Extract answer and file citations
        answer = ""
        references = []
        file_citations = set()  # To track unique files
        
        for item in response.output:
            if item.type == "message" and hasattr(item, "content"):
                for content_item in item.content:
                    if hasattr(content_item, "text"):
                        answer = content_item.text
                        
                        # Extract file citations
                        if hasattr(content_item, "annotations"):
                            for annotation in content_item.annotations:
                                if annotation.type == "file_citation":
                                    citation_data = {
                                        "file_id": annotation.file_id,
                                        "filename": annotation.filename
                                    }
                                    # Only add unique file citations
                                    citation_key = f"{annotation.file_id}:{annotation.filename}"
                                    if citation_key not in file_citations:
                                        file_citations.add(citation_key)
                                        references.append(format_reference(citation_data))
        
        no_info_phrases = [
            "Sorry",
            "I'm sorry",
            "don't know", 
            "couldn't find", 
            "no information", 
            "not mentioned", 
            "not available",
            "not in the context",
            "not found in the documents",
            "does not contain information",
        ]
        
        if any(phrase in answer.lower() for phrase in no_info_phrases):
            print("Response indicates no relevant information - returning without references")
            return ChatResponse(answer=answer, references=[])
            
        print(f"Returning answer with {len(references)} unique file references")
        return ChatResponse(answer=answer, references=references)
    
    except Exception as e:
        print(f"Error in text chat: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")

@app.post("/rtc-connect")
async def book_connect_rtc(request: Request):
    """Unified WebRTC endpoint with dynamic intent routing"""
    print("Unified RTC connection request received")
    
    # Common setup for both flows
    client_sdp = await request.body()
    if not client_sdp:
        raise HTTPException(status_code=400, detail="No SDP provided")
    
    client_sdp = client_sdp.decode()

    context_summary = []
    for doc_id, doc_info in document_metadata.items():
        content = doc_info.get("content", {})
        for page_num in sorted(content.keys())[:]:
            page_text = content[page_num].get("text", "")
            if page_text:
                preview = page_text[:330]
                context_summary.append(preview)
        
    context = "\n".join(context_summary)
        
    print(f"Providing document summary context of length: {len(context)}")
    
    instructions = f"""You are a versatile company support assistant that can handle both general support questions and booking requests. When answering general support questions, use the file search tool to retrieve information from company documents.

DETERMINE USER INTENT:
- When the user explicitly says they want to create a booking with phrases like: "I want to create a booking", "I need to book a shipment", "I want to send a package", "I need to arrange a pickup" - use the BOOKING WORKFLOW.
- When the user asks about HOW to create a booking with phrases like: "how do I create a booking", "how to book a shipment", "how to add a booking" - use the GENERAL SUPPORT workflow to explain the process using information from documents
- For all other queries - use the GENERAL SUPPORT workflow with document search.

GENERAL SUPPORT WORKFLOW:
1. Use only information from the company documents accessed through the file search tool
2. If you don't find relevant information, politely say you don't know
3. Keep responses concise and clear for voice interaction

BOOKING WORKFLOW:
1. Collect information in this exact order:
   - Full package sender address (street, city, country)
   - Full package receiver address (street, city, country)
   - Package or document description
   - package Weight in kilograms
   - package Dimensions (length × width × height in cm)
   - package Special handling requirements
2. Ask if package return pickup is needed (yes/no)
3. Ask one question at a time, moving to the next after receiving an answer
4. After collecting all information:
   - Generate a single 6-digit booking number
   - Display all collected details in this format:
     Booking number: [number]
     Sender address: [sender address]
     Receiver address: [receiver address]
     Package description: [description]
     Weight: [weight] kg
     Dimensions: [dimensions] cm
     Special handling: [requirements]
     Return pickup: [yes/no]
5. Ask for final confirmation before completing the booking
6. If the user wants to correct details, update only the specified part

Remember to stay conversational and helpful throughout the interaction."""
    
    try:
        async with httpx.AsyncClient() as client:
            # Get session token
            token_res = await client.post(
                OPENAI_SESSION_URL,
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
                json={
                    "model": MODEL_ID, 
                    "modalities": ["audio", "text"],
                    "voice": VOICE, 
                    "input_audio_format": "pcm16",
                    "output_audio_format": "pcm16",
                    "input_audio_transcription": {
                        "model": "whisper-1",
                        "language": "en"
                    },
                    "tools": [  # Add file search tool here
                        {
                            "type": "file_search",
                            "vector_store_ids": [VECTOR_STORE_ID]
                        }
                    ]
                }
            )

            if token_res.status_code != 200:
                raise HTTPException(status_code=500, detail="Token request failed")

            token_data = token_res.json()
            ephemeral_token = token_data.get('client_secret', {}).get('value', '')

            # Dynamic parameter configuration
            params = {
                "model": MODEL_ID,
                "instructions": instructions,
                "voice": VOICE,
            }

            # Single SDP exchange with dynamic instructions
            sdp_res = await client.post(
                OPENAI_API_URL,
                headers={
                    "Authorization": f"Bearer {ephemeral_token}",
                    "Content-Type": "application/sdp"
                },
                params=params,
                content=client_sdp
            )

            return Response(
                content=sdp_res.content,
                media_type='application/sdp',
                status_code=sdp_res.status_code
            )

    except Exception as e:
        return Response(
            status_code=500,
            content={"detail": f"Unified assistant error: {str(e)}"}
        )     
     
@app.post("/references_for_query", response_model=ChatResponse)
async def references_for_query(transcript: ChatTranscript):
    """Retrieve references for a given text query."""
    print(f"Query received: {transcript.question}")
    print(f"Response received: {transcript.answer}")
    
    try:
        # Simple greetings and thanks in query that don't need references
        simple_query_phrases = [
            "hii", "hi", "hello", "hey", "bye", "goodbye", "thanks", "thank you"
        ]
        
        if any(phrase == transcript.question.lower().strip() or 
               f" {phrase} " in f" {transcript.question.lower().strip()} " 
               for phrase in simple_query_phrases):
            print("Simple greeting query - returning without references")
            return ChatResponse(answer=transcript.question, references=[])
        
        # Complete phrases that indicate no relevant information in the answer
        no_info_phrases = [
            "don't know", 
            "couldn't find", 
            "no information", 
            "not mentioned", 
            "not available",
            "not in the context",
            "not found in the documents",
            "does not contain information"
        ]
        
        # Full phrases that should not trigger references
        exclude_full_phrases = [
            "booking confirmed", 
            "booking number:", 
            "booking details",
            "special handling", 
            "return pickup", 
            "package",
            "street, city, and country?",
            "length, width, and height?"
            "package description", 
            "kilograms?", 
            "centimeters?",
            "thank you",
            "summary",
            "welcome"
        ]
        
        # Check if any no_info_phrases are in the answer
        if any(phrase in transcript.answer.lower() for phrase in no_info_phrases):
            print("Response indicates no relevant information - returning without references")
            return ChatResponse(answer=transcript.question, references=[])
        
        # Check for complete phrases to exclude
        if any(f" {phrase} " in f" {transcript.answer.lower()} " for phrase in exclude_full_phrases):
            print("Response contains specific phrases that don't need references")
            return ChatResponse(answer=transcript.question, references=[])
        
        # Use OpenAI's file search tool to find references
        response = client.responses.create(
            model=CHAT_MODEL,
            input=transcript.question,
            tools=[{
                "type": "file_search",
                "vector_store_ids": [VECTOR_STORE_ID]
            }]
        )
        
        # Extract file citations
        references = []
        file_citations = set()  # To track unique files
        
        for item in response.output:
            if item.type == "message" and hasattr(item, "content"):
                for content_item in item.content:
                    if hasattr(content_item, "annotations"):
                        for annotation in content_item.annotations:
                            if annotation.type == "file_citation":
                                citation_data = {
                                    "file_id": annotation.file_id,
                                    "filename": annotation.filename
                                }
                                # Only add unique file citations
                                citation_key = f"{annotation.file_id}:{annotation.filename}"
                                if citation_key not in file_citations:
                                    file_citations.add(citation_key)
                                    references.append(format_reference(citation_data))
        
        if not references:
            print("No relevant documents found")
            return ChatResponse(answer=transcript.question, references=[])
        
        print(f"Found {len(references)} relevant documents")
        
        return ChatResponse(answer=transcript.question, references=references)
    
    except Exception as e:
        print(f"Error in references_for_query: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")

@app.on_event("startup")
async def startup_event():
    """Load existing document metadata on startup."""
    print("Starting application...")
    await load_existing_documents()
    print(f"Startup complete. Total documents loaded: {len(document_metadata)}")
    

if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8001, reload=True)